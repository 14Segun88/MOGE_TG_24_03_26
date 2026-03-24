"""
SverkaChecker — модуль сверки проектной документации по таблице Владимира.

Назначение:
  Читает DOCX-файл с таблицей сравнения ТЗ и ПЗ, извлекает требования,
  сопоставляет их с текстом из ПД-пакета.

Использование в пайплайне:
  checker = SverkaChecker("sverka_ot_vladimira.docx")
  result = checker.check(pd_text)

Поддерживает:
  - .docx (через python-docx)
  - .txt (простой текст)
  - .pdf (через fitz/PyMuPDF)
"""
from __future__ import annotations

import logging
import os
from dataclasses import dataclass, field
from typing import Optional

log = logging.getLogger("sverka_checker")


@dataclass
class SverkaItem:
    """Одно требование из таблицы сверки."""
    requirement: str      # Что должно быть (из ТЗ)
    expected: str = ""    # Ожидаемое значение/формулировка
    found_in_pd: str = "" # Что нашли в ПД
    compliant: bool | None = None   # True/False/None (не удалось проверить)
    comment: str = ""


@dataclass
class SverkaResult:
    """Итог сверки ПД по таблице Владимира."""
    source_file: str
    total_items: int = 0
    compliant_count: int = 0
    non_compliant_count: int = 0
    skipped_count: int = 0
    items: list[SverkaItem] = field(default_factory=list)
    error: str = ""

    @property
    def compliance_rate(self) -> float:
        if self.total_items == 0:
            return 0.0
        return self.compliant_count / self.total_items

    @property
    def is_compliant(self) -> bool:
        return self.non_compliant_count == 0 and self.total_items > 0


class SverkaChecker:
    """
    Читает таблицу сверки ТЗ/ПЗ из DOCX/TXT/PDF и
    сопоставляет требования с текстом проектной документации.
    """

    def __init__(self, sverka_path: str):
        self.path = sverka_path
        self._requirements: list[dict] = []
        self._loaded = False

    # ──────────────────────────────────────────
    #  Загрузка документа сверки
    # ──────────────────────────────────────────

    def load(self) -> bool:
        """Загрузить таблицу сверки из файла. Возвращает True при успехе."""
        if not os.path.exists(self.path):
            log.error(f"Файл сверки не найден: {self.path}")
            return False

        if os.path.getsize(self.path) == 0:
            log.error(f"Файл сверки пустой: {self.path}")
            return False

        ext = os.path.splitext(self.path)[1].lower()
        try:
            if ext == ".docx":
                self._load_docx()
            elif ext == ".txt":
                self._load_txt()
            elif ext == ".pdf":
                self._load_pdf()
            else:
                log.warning(f"Неизвестный формат: {ext}. Пробую как TXT.")
                self._load_txt()

            log.info(f"Loaded {len(self._requirements)} требований из: {self.path}")
            self._loaded = True
            return True
        except Exception as e:
            log.error(f"Ошибка загрузки сверки: {e}")
            return False

    def _load_docx(self):
        """Читает DOCX: таблицы (ТЗ | ПЗ) и параграфы с требованиями."""
        from docx import Document as DocxDocument
        doc = DocxDocument(self.path)

        # Приоритет 1: таблицы (формат ТЗ | ПЗ | Примечание)
        for tbl in doc.tables:
            header_row = [c.text.strip() for c in tbl.rows[0].cells] if tbl.rows else []
            
            # Определяем индексы нужных колонок
            tz_col = next((i for i, h in enumerate(header_row) 
                          if "ТЗ" in h or "требование" in h.lower() or "задание" in h.lower()), 0)
            pz_col = next((i for i, h in enumerate(header_row) 
                          if "ПЗ" in h or "проект" in h.lower()), 1 if len(header_row) > 1 else 0)

            for row in tbl.rows[1:]:  # Пропускаем заголовок
                cells = row.cells
                if len(cells) < 2:
                    continue
                tz_text = cells[tz_col].text.strip() if tz_col < len(cells) else ""
                pz_text = cells[pz_col].text.strip() if pz_col < len(cells) else ""
                if tz_text:
                    self._requirements.append({
                        "requirement": tz_text,
                        "expected": pz_text,
                        "source": "table"
                    })

        # Приоритет 2: параграфы с маркерами требований (если таблиц нет)
        if not self._requirements:
            markers = ["требован", "должен", "необходим", "предусмотрен", "обеспечен"]
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if any(m in text.lower() for m in markers):
                    self._requirements.append({
                        "requirement": text,
                        "expected": "",
                        "source": "paragraph"
                    })

    def _load_txt(self):
        """Читает TXT: каждая непустая строка — отдельное требование."""
        with open(self.path, encoding="utf-8", errors="replace") as f:
            lines = f.readlines()
        for line in lines:
            line = line.strip()
            if len(line) > 15:  # Пропускаем очень короткие строки
                self._requirements.append({
                    "requirement": line,
                    "expected": "",
                    "source": "text"
                })

    def _load_pdf(self):
        """Читает PDF через PyMuPDF (fitz)."""
        import fitz
        text_chunks = []
        with fitz.open(self.path) as doc:
            for page in doc:
                text_chunks.append(page.get_text())
        full_text = "\n".join(text_chunks)
        
        markers = ["требован", "должен", "необходим", "предусмотрен"]
        for line in full_text.split("\n"):
            line = line.strip()
            if len(line) > 20 and any(m in line.lower() for m in markers):
                self._requirements.append({
                    "requirement": line,
                    "expected": "",
                    "source": "pdf"
                })

    # ──────────────────────────────────────────
    #  Сверка с текстом ПД
    # ──────────────────────────────────────────

    def check(self, pd_text: str, use_llm: bool = False) -> SverkaResult:
        """
        Сверяет требования таблицы с текстом ПД.
        
        Args:
            pd_text: Текст из проектной документации
            use_llm: Если True — вызывает LLM для сложных случаев (медленнее)
        """
        if not self._loaded:
            self.load()

        result = SverkaResult(source_file=self.path)

        if not self._requirements:
            result.error = "Не удалось загрузить требования из файла сверки"
            return result

        pd_text_lower = pd_text.lower()
        result.total_items = len(self._requirements)

        for req_dict in self._requirements:
            req_text = req_dict["requirement"]
            expected = req_dict.get("expected", "")

            item = SverkaItem(requirement=req_text, expected=expected)

            # Простая текстовая проверка: ищем ключевые слова из требования в ПД
            keywords = self._extract_keywords(req_text)
            found_keywords = [k for k in keywords if k in pd_text_lower]
            
            if len(found_keywords) >= max(1, len(keywords) // 2):
                # Нашли большинство ключевых слов
                item.compliant = True
                item.found_in_pd = ", ".join(found_keywords[:5])
                item.comment = f"Найдено {len(found_keywords)}/{len(keywords)} ключевых слов"
                result.compliant_count += 1
            elif not keywords:
                item.compliant = None
                item.comment = "Не удалось извлечь ключевые слова"
                result.skipped_count += 1
            else:
                item.compliant = False
                item.comment = f"Найдено {len(found_keywords)}/{len(keywords)} ключевых слов: {', '.join(found_keywords[:3])}"
                result.non_compliant_count += 1

            result.items.append(item)

        log.info(
            f"Сверка завершена: {result.compliant_count}/{result.total_items} "
            f"соответствуют, {result.non_compliant_count} нарушений"
        )
        return result

    def _extract_keywords(self, text: str) -> list[str]:
        """Извлекает ключевые слова из требования (фильтруем стоп-слова)."""
        stop_words = {
            "должен", "должна", "должно", "должны", "быть", "является", "являются",
            "при", "для", "или", "если", "что", "как", "это", "также", "все",
            "следует", "необходимо", "требуется", "обеспечивает", "предусматривается",
            "на", "в", "с", "к", "о", "и", "а", "но", "по", "из", "от", "до",
        }
        words = []
        for word in text.lower().split():
            word = word.strip(".,;:()\"-'«»")
            if len(word) > 4 and word not in stop_words:
                words.append(word)
        return words[:10]  # Топ-10 ключевых слов


# ──────────────────────────────────────────────────────────────────────────────
#  Быстрый тест
# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO)

    docx_path = sys.argv[1] if len(sys.argv) > 1 else (
        "АВТОМАТИЗИРОВАННАЯ СИСТЕМА ЭКСПЕРТИЗЫ ДОКУМЕНТАЦИИ.docx"
    )
    
    checker = SverkaChecker(docx_path)
    checker.load()
    
    print(f"\\n--- Загружено {len(checker._requirements)} требований ---")
    for i, r in enumerate(checker._requirements[:5], 1):
        print(f"{i}. [{r['source']}] {r['requirement'][:100]}")
    
    print("\\n--- Тестовая сверка ---")
    test_pd = (
        "Проект разработан в соответствии с требованиями ПП РФ №963. "
        "Раздел 1 (ПЗ) содержит технические характеристики здания. "
        "Пожарная безопасность обеспечена по нормам СП 1.13130. "
        "Площадь здания 1500 кв.м, этажность 5 этажей."
    )
    result = checker.check(test_pd)
    print(f"Соответствует: {result.compliant_count}/{result.total_items}")
    print(f"Нарушений: {result.non_compliant_count}")
    print(f"Доля соответствия: {result.compliance_rate:.0%}")
