#!/usr/bin/env python3
"""
Утилита для чтения DOCX-файлов без иероглифов.

Использование:
  python3 read_docx.py <путь/к/файлу.docx>
  python3 read_docx.py sverka_ot_vladimira.docx

Проблема с иероглифами:
  Файлы .docx — это ZIP-архивы с XML-файлами внутри.
  VS Code / текстовые редакторы открывают их как бинарный файл → иероглифы.
  Этот скрипт распаковывает XML и выводит чистый текст через python-docx.
"""
import sys
import os
import importlib.util


def check_docx_library():
    """Проверяем наличие python-docx."""
    spec = importlib.util.find_spec("docx")
    if spec is None:
        print("❌ Библиотека python-docx не установлена.")
        print("   Установите: .venv/bin/pip install python-docx")
        sys.exit(1)


def read_docx_text(path: str) -> str:
    """Читает DOCX файл и возвращает чистый текст."""
    from docx import Document
    doc = Document(path)
    
    lines = []
    
    # Параграфы
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    
    # Таблицы
    for tbl_idx, tbl in enumerate(doc.tables):
        lines.append(f"\n=== Таблица {tbl_idx + 1} ===")
        for row in tbl.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)
    
    return "\n".join(lines)


def main():
    check_docx_library()
    
    if len(sys.argv) < 2:
        # Если путь не указан — ищем первый .docx в текущей папке
        docx_files = [f for f in os.listdir(".") if f.endswith(".docx")]
        if not docx_files:
            print("Использование: python3 read_docx.py <файл.docx>")
            sys.exit(1)
        path = docx_files[0]
        print(f"Автовыбор: {path}")
    else:
        path = sys.argv[1]
    
    if not os.path.exists(path):
        print(f"❌ Файл не найден: {path}")
        sys.exit(1)
    
    if os.path.getsize(path) == 0:
        print(f"❌ Файл пустой (0 байт): {path}")
        print("   Скопируйте нужный DOCX с Windows в эту папку:")
        print("   cp '/mnt/c/Users/123/Downloads/Telegram Desktop/Таблица_сравнения.docx' .")
        sys.exit(1)
    
    print(f"📄 Читаем: {path}")
    print("=" * 60)
    
    text = read_docx_text(path)
    print(text)
    
    # Сохраняем в .txt для удобного просмотра
    txt_path = os.path.splitext(path)[0] + "_extracted.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"\n✅ Текст сохранён в: {txt_path}")
    print(f"   Строк: {len(text.splitlines())}, символов: {len(text):,}")


if __name__ == "__main__":
    main()
